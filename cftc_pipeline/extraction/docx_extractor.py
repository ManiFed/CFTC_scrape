"""Word document text extraction."""
from __future__ import annotations

import io
import logging
from typing import Optional

from cftc_pipeline.extraction.pdf_extractor import ExtractionOutput

logger = logging.getLogger(__name__)


def extract_docx(data: bytes) -> ExtractionOutput:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        full_text = "\n\n".join(paragraphs)
        status = "ok" if full_text.strip() else "empty"
        return ExtractionOutput(
            text=full_text,
            page_count=len(doc.paragraphs),
            method="python-docx",
            status=status,
        )
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ExtractionOutput(
            text="", page_count=0, method="python-docx", status="failed", error=str(exc)
        )
